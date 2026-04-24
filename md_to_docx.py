#!/usr/bin/env python3
"""
Convert a markdown cover letter to a clean Word .docx.

Minimal markdown handling:
- `#` / `##` headers → dropped
- Blank line → paragraph break
- Trailing two spaces (`  `) at end of line → HARD LINE BREAK within a paragraph
  (useful for signature blocks and address headers)
- `**bold**` → bold runs
- `*italic*` → italic runs

Body font: Calibri 11pt, 1-inch margins, 6pt space after paragraphs.

Usage:
    python3 md_to_docx.py input.md output.docx
"""
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
ITALIC_RE = re.compile(r"(?<!\*)\*(.+?)\*(?!\*)")


def render_inline(paragraph, text):
    """Add runs to paragraph, preserving **bold** / *italic* inline."""
    remaining = text
    while remaining:
        bold_match = BOLD_RE.search(remaining)
        italic_match = ITALIC_RE.search(remaining)
        candidates = [m for m in [bold_match, italic_match] if m]
        if not candidates:
            run = paragraph.add_run(remaining)
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            break
        first = min(candidates, key=lambda m: m.start())
        before = remaining[: first.start()]
        if before:
            run = paragraph.add_run(before)
            run.font.name = "Calibri"
            run.font.size = Pt(11)
        run = paragraph.add_run(first.group(1))
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        if first is bold_match:
            run.bold = True
        else:
            run.italic = True
        remaining = remaining[first.end():]


def add_paragraph(doc, segments):
    """Add one paragraph from a list of (text, hard_break_before_next) segments."""
    if not segments:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    for i, (text, hard_break_after) in enumerate(segments):
        render_inline(p, text)
        if hard_break_after and i < len(segments) - 1:
            br_run = p.add_run()
            br_run.add_break()


def convert(md_path: Path, docx_path: Path):
    md_text = md_path.read_text()
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md_text.splitlines()

    # Build list of paragraph segments. Each paragraph is a list of (text, hard_break_after).
    current_segments = []

    def flush():
        if current_segments:
            # Drop the hard_break from the last segment (no break after paragraph end)
            current_segments[-1] = (current_segments[-1][0], False)
            add_paragraph(doc, current_segments)
            current_segments.clear()

    for raw_line in lines:
        # Detect hard break (trailing 2+ spaces on a non-blank line)
        hard_break = raw_line.endswith("  ")
        line = raw_line.rstrip()

        # Markdown headers — skip
        if line.lstrip().startswith("#"):
            flush()
            continue
        # Horizontal rule — skip
        if line.strip() in ("---", "***", "___"):
            flush()
            continue

        if not line.strip():
            flush()
            continue

        current_segments.append((line, hard_break))

    flush()
    doc.save(str(docx_path))


def main():
    if len(sys.argv) != 3:
        sys.exit("Usage: python3 md_to_docx.py input.md output.docx")
    convert(Path(sys.argv[1]), Path(sys.argv[2]))
    print(f"✅ {sys.argv[2]}")


if __name__ == "__main__":
    main()
