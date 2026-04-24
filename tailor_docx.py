#!/usr/bin/env python3
"""
Tailor a master .docx resume in-place to a specific JD while preserving the
original layout (fonts, spacing, indentation, bold styling).

Invoked by per-job build scripts that import `tailor()`.

Configuration (read from environment with fallbacks):
    RESUME_SRC_DOCX   Path to the candidate's master .docx (required)

CONTRACT of the source docx (must stay stable — tune JOB_RANGES to match):
    Para 0: Name
    Para 1: Contact line
    Para 3: "Professional Summary:" header
    Para 5: Summary paragraph
    Para 7: "Technical Skills:" header
    Paras 8-15: Skill category lines
    Para 16: "Professional Experience:" header
    Para 17+: Per-job header + bullets (see JOB_RANGES)
"""
import os
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn


def _resolve_source_docx() -> Path:
    env_val = os.environ.get("RESUME_SRC_DOCX")
    if env_val:
        p = Path(os.path.expanduser(env_val))
        if p.is_file():
            return p
    # Fallback: config/profile.yml `resume.source_docx`
    try:
        import yaml
        cfg_path = Path(__file__).parent / "config" / "profile.yml"
        if cfg_path.exists():
            cfg = yaml.safe_load(cfg_path.read_text()) or {}
            cand = cfg.get("resume", {}).get("source_docx")
            if cand:
                return Path(os.path.expanduser(cand))
    except Exception:
        pass
    raise SystemExit(
        "❌ Cannot locate master resume .docx.\n"
        "   Set RESUME_SRC_DOCX env variable OR add `resume.source_docx: <path>` to config/profile.yml."
    )


SRC = _resolve_source_docx()

# Bullet index ranges in source
JOB_RANGES = {
    "fifth_third": (18, 27),   # 10 bullets
    "optum":       (29, 35),   # 7 bullets
    "goldman":     (38, 43),   # 6 bullets
    "gainwell":    (47, 51),   # 5 bullets
    "pitney":      (55, 59),   # 5 bullets
}
SUMMARY_IDX = 5


def _replace_paragraph_text(paragraph, new_text):
    """Replace paragraph text while preserving the formatting of the first run."""
    # Keep the first run, set its text to the full new text, wipe the others.
    if not paragraph.runs:
        paragraph.text = new_text
        return
    first_run = paragraph.runs[0]
    first_run.text = new_text
    # Remove remaining runs
    for run in paragraph.runs[1:]:
        run._element.getparent().remove(run._element)


def _delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def _clone_paragraph_after(template_paragraph, new_text):
    """Insert a new paragraph with cloned formatting right after template_paragraph, containing new_text."""
    new_p = deepcopy(template_paragraph._element)
    template_paragraph._element.addnext(new_p)
    # Now wipe text and set new text on the first run
    # Find all w:t elements and clear them
    for t in new_p.findall(qn('w:r')):
        new_p.remove(t)
    # Copy a run element template
    template_run = template_paragraph._element.findall(qn('w:r'))[0]
    fresh_run = deepcopy(template_run)
    # Clear any text inside fresh_run and set new
    for t_el in fresh_run.findall(qn('w:t')):
        fresh_run.remove(t_el)
    # Create a new w:t element
    from docx.oxml import OxmlElement
    t_el = OxmlElement('w:t')
    t_el.text = new_text
    t_el.set(qn('xml:space'), 'preserve')
    fresh_run.append(t_el)
    new_p.append(fresh_run)
    return new_p


def _set_job_bullets(doc, job_key, new_bullets):
    """Replace the bullets for `job_key` with `new_bullets` (list of strings).

    If new_bullets has fewer items than the original, delete extras.
    If new_bullets has more, clone the last bullet and append.
    """
    start, end = JOB_RANGES[job_key]
    # Work on live paragraphs (list indices change as we delete, so take snapshot)
    paras = doc.paragraphs
    original_bullets = paras[start:end+1]
    original_count = len(original_bullets)
    new_count = len(new_bullets)

    # Replace first min(new_count, original_count) bullets
    for i in range(min(new_count, original_count)):
        _replace_paragraph_text(original_bullets[i], new_bullets[i])

    # Delete any extras if new < original
    if new_count < original_count:
        for p in original_bullets[new_count:]:
            _delete_paragraph(p)

    # Append extras by cloning if new > original
    if new_count > original_count:
        template = original_bullets[-1]
        # Insert after the last bullet, cloning formatting
        for extra_text in new_bullets[original_count:]:
            new_p_el = _clone_paragraph_after(template, extra_text)
            # Update template reference so subsequent clones go in the right order
            class _Wrapper:
                def __init__(self, el): self._element = el
            template = _Wrapper(new_p_el)


from typing import Union


def _candidate_slug() -> str:
    """Read candidate name from config/profile.yml and slugify (kebab-case)."""
    try:
        import yaml
        cfg_path = Path(__file__).parent / "config" / "profile.yml"
        if cfg_path.exists():
            cfg = yaml.safe_load(cfg_path.read_text()) or {}
            name = cfg.get("candidate", {}).get("full_name", "")
            if name:
                return "-".join(name.lower().split())
    except Exception:
        pass
    return "resume"


def tailor(
    company_slug: str,
    out_dir: Path,
    summary: str,
    bullets: dict,           # {"fifth_third": [...], "optum": [...], ...}
    skills_override: Union[dict, None] = None,   # optional {category: new_line}
    location_override: Union[str, None] = None,
):
    """Produce a tailored .docx at out_dir / <candidate_slug>.docx + .pdf per company."""
    from datetime import date
    doc = Document(str(SRC))

    # Recompute indices after mutations (we walk sections in source order)
    # Strategy: mutate jobs from bottom up so earlier indices stay valid.
    order = ["pitney", "gainwell", "goldman", "optum", "fifth_third"]

    # Summary first (no index shift yet)
    _replace_paragraph_text(doc.paragraphs[SUMMARY_IDX], summary)

    # Skills overrides: simple find-and-replace at line level
    if skills_override:
        for p in doc.paragraphs[7:16]:
            for prefix, new_line in skills_override.items():
                if p.text.startswith(prefix):
                    _replace_paragraph_text(p, new_line)

    # Now apply bullet changes from last to first so ranges stay valid
    for job_key in order:
        if job_key in bullets:
            _set_job_bullets(doc, job_key, bullets[job_key])

    # Save docx into per-company subfolder: output/{company_slug}/
    # Filename uses candidate slug derived from config/profile.yml full_name
    # (or "resume" fallback). Company context is the folder.
    company_dir = out_dir / company_slug
    company_dir.mkdir(parents=True, exist_ok=True)
    slug = _candidate_slug()
    docx_path = company_dir / f"{slug}.docx"
    pdf_path = company_dir / f"{slug}.pdf"
    doc.save(str(docx_path))
    print(f"✅ docx written: {docx_path}")

    # Convert to PDF via Microsoft Word (AppleScript)
    import subprocess
    applescript = f'''tell application "Microsoft Word"
  set docxPath to POSIX file "{docx_path}" as alias
  open docxPath
  set docRef to active document
  save as docRef file name "{pdf_path}" file format format PDF
  close docRef saving no
end tell'''
    result = subprocess.run(
        ["osascript", "-e", applescript],
        capture_output=True, text=True, timeout=120,
    )
    if pdf_path.exists():
        size_kb = pdf_path.stat().st_size / 1024
        # Authoritative page count via pypdf (mdls caches; pypdf reads the file)
        try:
            from pypdf import PdfReader
            pages = len(PdfReader(str(pdf_path)).pages)
        except ImportError:
            pages_result = subprocess.run(
                ["mdls", "-name", "kMDItemNumberOfPages", str(pdf_path)],
                capture_output=True, text=True,
            )
            pages = int(pages_result.stdout.strip().split("=")[-1].strip() or 0)
        status = "✅" if pages <= 2 else "❌"
        print(f"{status} pdf written: {pdf_path} ({size_kb:.0f} KB, {pages} pages)")
        if pages > 2:
            raise SystemExit(
                f"\n🚨 HARD STOP — output is {pages} pages, must be ≤ 2.\n"
                f"   Trim the longest bullets in build_<company>_docx.py (keep bullet count, shorten text).\n"
                f"   DO NOT deliver this file."
            )
    else:
        print(f"⚠️ PDF conversion failed. stderr: {result.stderr}")
        print("   Open the .docx manually in Word and File → Save As → PDF.")

    return docx_path, pdf_path


if __name__ == "__main__":
    import sys
    print("This module is meant to be imported. Use build_<company>_docx.py scripts.")
    sys.exit(1)
